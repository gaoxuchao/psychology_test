import subject_db
from test_database import subject
from test_database import pd_columns
import pandas as pd
import numpy as np
from docx import Document
from enum import Enum

class test_columns(Enum):
    TYPE       = "题目类型"
    CONTENT    = "内容" 
    PAGE_NUM   = "页码"
    IMPORTANCE = "重要性"
    RIGHT_NUM  = "答对次数"
    WRONG_NUM  = "答错次数"
    TEST_CNT   = "出题次数"
    SUB_NAME   = "科目"
    SUB_ID     = "科目编号"


class test_main() :

    def __init__(self):

        self.test_base_path = r".\试卷库" "\\"
        self.subject_list = [] # 创建空列表

        subject_id = 0
        for sub in subject: # 在枚举类型中遍历
            # 初始化各个科目的对象
            self.subject_list.append(subject_db.subject_db(sub,subject_id))
            subject_id +=1

        for psy_sub in self.subject_list:
            # 获取各个科目的数据
            psy_sub.get_data_form_excel(psy_sub.sub_xlsx_file_name + '.xlsx')

        self.test_db_init()


    def test_db_init(self):
        '''
            初始化题库三种题型的database内容，只保留初始的列名称
        '''
        # 创建pandas对象
        # 根据列标题创建列表字典
        dic_tmp = {}
        for title in test_columns:
            dic_tmp[title.value] = np.nan
        pd_tmp = [dic_tmp]

        self.des_set = pd.DataFrame(pd_tmp)
        self.sim_des_set = pd.DataFrame(pd_tmp)
        self.term_expl_set = pd.DataFrame(pd_tmp)

        self.test_db_clear() #调用database清空函数

    def test_db_clear(self):
        self.des_set.drop(self.des_set.index,inplace=True)
        self.sim_des_set.drop(self.sim_des_set.index,inplace=True)
        self.term_expl_set.drop(self.term_expl_set.index,inplace=True)

    def init_xlxs_from_docs(self):
        '''
        从原始WORD文档中，重置所有科目的EXCEL题库
        '''
        for psy_sub in self.subject_list:
            psy_sub.parser_subject()

    def make_test_db(self,mark_db:bool=True):
        '''
        从各科目题库中随机抽题，存入数据库中
            mark_db : 出题时是否反标题库，默认打开，测试模式下可关闭
        '''       

        self.test_db_clear()
        
        for i in range(len(self.subject_list)):
            if (i == 0): #init test 
                self.des_set      = self.subject_list[i].auto_quiz_setter_des()
                self.sim_des_set  = self.subject_list[i].auto_quiz_setter_sim_des()
                self.term_expl_set= self.subject_list[i].auto_quiz_setter_term_expl()
                self.des_set['科目']     = self.subject_list[i].subject_name
                self.des_set['科目编号'] = self.subject_list[i].subject_id 
                self.sim_des_set['科目']     = self.subject_list[i].subject_name
                self.sim_des_set['科目编号'] = self.subject_list[i].subject_id 
                self.term_expl_set['科目']     = self.subject_list[i].subject_name
                self.term_expl_set['科目编号'] = self.subject_list[i].subject_id 

            else:
                tmp_dp0 = self.subject_list[i].auto_quiz_setter_des()
                tmp_dp1 = self.subject_list[i].auto_quiz_setter_sim_des()
                tmp_dp2 = self.subject_list[i].auto_quiz_setter_term_expl() 
                tmp_dp0['科目']     = self.subject_list[i].subject_name
                tmp_dp0['科目编号'] = self.subject_list[i].subject_id 
                tmp_dp1['科目']     = self.subject_list[i].subject_name
                tmp_dp1['科目编号'] = self.subject_list[i].subject_id 
                tmp_dp2['科目']     = self.subject_list[i].subject_name
                tmp_dp2['科目编号'] = self.subject_list[i].subject_id 
                self.des_set      = self.des_set.append      ( tmp_dp0)
                self.sim_des_set  = self.sim_des_set.append  ( tmp_dp1)
                self.term_expl_set= self.term_expl_set.append( tmp_dp2) 
            
            #处理速度优化，每抽完一个科目，保存一次EXCEL文件
            self.subject_list[i].save_database_to_excel(self.subject_list[i].sub_xlsx_file_name + '.xlsx',
                                                        'sheet1')     



    def create_test_doc(self,test_no:int):
        '''
            将随机抽取后的题库输出到Word文档中
                test_no:试卷编号
        '''

        test_doc = Document()
        test_name = "试卷" + str(test_no)

        test_doc.add_heading(test_name)

        test_doc.add_heading("名词解释 (3分)",level=1)
        for i in range(len(self.term_expl_set.index)):
            no_tmp = i +1
            dp_tmp = self.term_expl_set.iloc[i]
            content_tmp = dp_tmp["内容"]
            page_num_tmp = dp_tmp["页码"]
            para_tmp = str(no_tmp) +". " + content_tmp + "【" + str(page_num_tmp) + "】"
            p_term = test_doc.add_paragraph(para_tmp)

        test_doc.add_heading("简答 (15分)",level=1)
        for i in range(len(self.sim_des_set.index)):
            no_tmp = i +1
            dp_tmp = self.sim_des_set.iloc[i]
            content_tmp = dp_tmp["内容"]
            page_num_tmp = dp_tmp["页码"]
            para_tmp = str(no_tmp) +". " + content_tmp + "【" + str(page_num_tmp) + "】"
            p_term = test_doc.add_paragraph(para_tmp)

        test_doc.add_heading("论述 (25分)",level=1)
        for i in range(len(self.des_set.index)):
            no_tmp = i +1
            dp_tmp = self.des_set.iloc[i]
            content_tmp = dp_tmp["内容"]
            page_num_tmp = dp_tmp["页码"]
            para_tmp = str(no_tmp) +". " + content_tmp + "【" + str(page_num_tmp) + "】"
            p_term = test_doc.add_paragraph(para_tmp)

        test_doc.add_page_break()

        test_doc.save(self.test_base_path + test_name + ".docx")


    def save_test_to_xlsx(self,test_no:int):
        '''
        保存test_db到试卷库路径下的EXCEL文件中
        ，将题号转换为单独的列
        ，为test_db新建索引
        ，试卷库路径：test_base_path
        ，保存的文件名为：“试卷+test_no+.xlsx”

        '''
        frames = [self.des_set,self.sim_des_set,self.term_expl_set]
        
        tmp_test_db = pd.concat(frames)

        tmp_test_db["题目索引"] = tmp_test_db.index

        #正确栏，默认为0，正确填1，错误填2
        tmp_test_db["正确"] = 0 

        # del tmp_test_db['内容']

        tmp_test_db["index"] = range(len(tmp_test_db))
        tmp_test_db = tmp_test_db.set_index(['index'])
        
        file_name =self.test_base_path + "试卷" + str(test_no) + '.xlsx'
        
        with pd.ExcelWriter(file_name) as writer: # pylint: disable=abstract-class-instantiated
            tmp_test_db.to_excel(writer,'sheet',index=True,index_label="index") #将pd写入example_file，索引列名为题号


    def make_test(self,test_no:int):
        '''
        创建试卷
            test_no : 试卷编号，必须为整数    
        '''
        self.make_test_db()
        self.save_test_to_xlsx(test_no)
        self.create_test_doc(test_no)

    def get_test_from_excel(self,test_no:int):
        '''
        从EXCEL中获取之前出过的试卷
            test_no: 试卷编号（必须在试卷库目录中存在相应的EXCEL文件）
            试卷库路径：test_base_path
            文件名为：“试卷+test_no+.xlsx”
        '''
        # 不要和test_db冲突
        # 
        file_name = self.test_base_path + "试卷" + str(test_no) + '.xlsx'
        with pd.ExcelFile(file_name) as xls:
            sheet_names = xls.sheet_names  #获取excel文件的sheet列表
            self.test_db_old = pd.read_excel(xls,sheet_names[0],index_col=0)  #指定第一列为索引列      


    def test_feedback(self,test_no:int):
        '''
        反馈答题结果：
           从EXCEL中获取答题结果，然后根据科目，科目编号和题目索引定位到具体题目
           根据正确与否，对答对次数和答错次数进行修改
           “正确”栏，默认为0，正确填1，错误填2
        '''
        self.get_test_from_excel(test_no)
        pass



                     
if __name__ == "__main__":
    print("######################################################")
    test = test_main()
    # test.init_xlxs_from_docs()
    test.make_test(2)
    # test.save_test_to_xlsx(0)
    # test.create_test_doc(4)
    print("######################################################")
