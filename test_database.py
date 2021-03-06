import re
from docx import Document
import xlwings as xw
import numpy as np
import pandas as pd
import re
from enum import Enum

class subject(Enum):
    COMM_PSY = '普通心理学'
    DVP_PSY  = '发展心理学'
    EDU_PSY  = '教育心理学'
    SOC_PSY  = '社会心理学'
    STA_PSY  = '心理统计学'
    MES_PSY  = '心理测量学'
    EXP_PSY  = '实验心理学'

class pd_columns(Enum): # no used
    TYPE       = "题目类型"
    CONTENT    = "内容" 
    PAGE_NUM   = "页码"
    IMPORTANCE = "重要性"
    RIGHT_NUM  = "答对次数"
    WRONG_NUM  = "答错次数"
    TEST_CNT   = "出题次数"


class database_proc(object):
    def __init__(self):

        #初始化pandas列标题
        self.columns_name = ["题目类型",
                            "内容",
                            "页码",
                            "重要性",
                            "答对次数",
                            "答错次数",
                            "出题次数"]

        self.database_init()

        #默认参数，用来初始化题目的默认重要性
        self.default_value = 1

        self.most_important = 3
        self.important      = 2
        self.no_important   = 1
        #print("父类初始化完成")


    def doc_parser(self,file_name):
        '''
           从word文档中获取数据，并做分类处理
           file_name : word文档路径+文件名
        '''
        page_num_pattern = r'\d{3}'               #页码正则匹配pattern，数字连续出现三次视为页码
        text_pattern = r'\w+'                     #文本匹配，至少要有一个文字
        page_num_re = re.compile(page_num_pattern)
        text_re     = re.compile(text_pattern)

        try :
            test_data = Document(file_name)
            for pgs in test_data.paragraphs:

                search_rslt = page_num_re.search(pgs.text)      #匹配页码
                text_rslt   = text_re.search(pgs.text)

                if search_rslt != None:
                    test_text   = page_num_re.sub('',pgs.text) #如果存在页码，将页码从文本中剔除
                    page_num    = search_rslt.group()          #获得页码
                else:
                    test_text   = pgs.text
                    #page_num 不变
                    #page_num    = np.nan
                
                if ("简答：" in pgs.text) or ("简述" in pgs.text):
                    test_type = '简答'
                elif "论述：" in pgs.text:
                    test_type = '论述'
                elif text_rslt != None:        #没有论述和简答标示，但是有其它文本，同意归类为名词解释
                    test_type = '名词解释'
                else:
                    test_type = None 

                if test_type != None:
                    #columns_name = ["题目类型","内容","页码","重要性","答对次数","答错次数"]
                    tmp_dict = {self.columns_name[0]:test_type,
                                self.columns_name[1]:test_text,
                                self.columns_name[2]:int(page_num),
                                self.columns_name[3]:self.default_value,
                                self.columns_name[4]:0,
                                self.columns_name[5]:0,
                                self.columns_name[6]:0}

                    self.database = self.database.append(tmp_dict,ignore_index=True)

        except OSError as os_e:
            print("文件打开异常！",os_e)
        finally:
            print(file_name,":搜索完成！")
        #endfunction:doc_parser        

    def doc_parser_comm(self,file_name,subject_name):
        '''
            将公共文档中的题目按照科目提取
            file_name : 文档路径加名称
            subject_name :科目名词，必须为字符串
        '''
        # test_no_pattern = r'\d{3}'                 #题号正则匹配pattern，数字连续出现三次视为页码
        # test_no_re = re.compile(test_no_pattern)
        subject_pattern = r'第\w+编 ' + subject_name
        subject_re      = re.compile(subject_pattern)
        subject_title   = r'第\w+编 '
        title_re        = re.compile(subject_title)
        para_found      = 0

        text_pattern    = r'\w+'
        text_re         = re.compile(text_pattern)

        importance_pattern = r'★{1,3}'
        importance_re      = re.compile(importance_pattern)

        importance      = 0
        test_type       = '简答'
        try :
            test_data = Document(file_name)
            for pgs in test_data.paragraphs:

                # search_rslt = test_no_re.search(pgs.text)      #匹配页码
                subject_search = subject_re.search(pgs.text)
                title_search   = title_re.search(pgs.text)

                if subject_search != None and para_found == 0:
                    para_found = 1
                    print('找到 '+pgs.text)
                    continue
                elif title_search != None and para_found ==1 :
                    para_found = 0
                    print(subject_name + "加载结束！")
                    # print(pgs.text)
                    break
                
                text_search = text_re.search(pgs.text)

                if para_found == 1 and text_search != None:
                    if "★★★" in pgs.text:
                        importance = self.most_important
                    elif "★★" in pgs.text:
                        importance =self.important
                    else:
                        importance = self.no_important

                    test_text = importance_re.sub("",pgs.text)

                    # columns_name = ["题目类型","内容","页码","重要性","答对次数","答错次数","出题次数"]
                    tmp_dict = {self.columns_name[0]:test_type,
                                self.columns_name[1]:test_text,
                                self.columns_name[2]:0,
                                self.columns_name[3]:importance,
                                self.columns_name[4]:0,
                                self.columns_name[5]:0,
                                self.columns_name[6]:0}

                    self.database = self.database.append(tmp_dict,ignore_index=True)

        except OSError as os_e:
            print("文件打开异常！",os_e)
        finally:
            print(file_name,' ：扫描完成！')
    

    def database_init(self):
        '''
            初始化database内容，只保留初始的列名称
        '''
        # 创建pandas对象
        # 根据列标题创建列表字典
        dic_tmp = {}
        for title in self.columns_name:
            dic_tmp[title] = np.nan
        pd_tmp = [dic_tmp]
        self.database =pd.DataFrame(pd_tmp)
        self.database_clear() #调用database清空函数

    def database_clear(self):
        '''
            清空database内容，只保留列名称
        '''
        self.database.drop(self.database.index,inplace=True)
    
    def get_data_form_excel(self,file_name,sheet_index = 0):
        '''
        从指定的文件中更新pandas数据库
        file_name : 指定的文件
        sheet_index：如果有多个sheet，可以指定第几个sheet，默认获取第1个sheet
        '''
        with pd.ExcelFile(file_name) as xls:
            self.database_sheet_names = xls.sheet_names  #获取excel文件的sheet列表
            self.database = pd.read_excel(xls,self.database_sheet_names[sheet_index],index_col=0)  #指定第一列为索引列 
            #FIX_BUG,在读取题库时，将题目类型中的空格全部清除  
            self.database["题目类型"] = self.database["题目类型"].str.replace(" ","")   

    def save_database_to_excel(self,file_name,sheet_name):
        '''
           将database写入指定EXCEL文件，同时必须指定sheet名称
           file_name : 指定的文件
           sheet_name : 写入的sheet名称
        '''
        with pd.ExcelWriter(file_name) as writer: # pylint: disable=abstract-class-instantiated
            self.database.to_excel(writer,sheet_name,index=True,index_label="题号") #将pd写入example_file，索引列名为题号
            # print(self.database.head(3))                #打印前3行元素
            # print(self.database.index)                  #得到dataframe的行元素
            # print(self.database.columns)                #得到dataframe的列元素
        
        app=xw.App(visible=False,add_book=False)  
        app.display_alerts=False
        app.screen_updating=False

        wb  = app.books.open(file_name)
        sht = wb.sheets[0]
        sht.autofit('c')                        #列宽自适应
        sht.range('a1:g1').color=(255,0,0)      #名称行设置为红色
        sht.range('c1').column_width = 80       #列宽设置为5
        # sht.autofit('r')   
        wb.save()
        wb.close()
        app.quit()
        


# wb = xw.Book(".\example.xlsx")

# # wb.save(".\example.xlsx")

# sht = wb.sheets("sheet1")

# sht.range('a6').expand('table').value = [['a','b','c'],['d','e','f'],['g','h','i']]

    # self.database.loc[2] = [2,'论述','ddd',345,3,0,0]
    # self.database.loc[3] = [3,'论述','aaa',354,3,0,0]    
    # print(self.database.loc[2,'题目内容'])        #取第3行的‘题目内容’列
    # print(self.database.loc[2,self.database.columns[3]]) #取第3行,4列的元素

if __name__ == "__main__":
    dp = database_proc()
    # dp.database_clear()

    print("######################################################")
    print(dp.database)
    print("######################################################")

    # dp.get_data_form_excel('example.xlsx')

    # dp.doc_parser_comm(r".\题库\心理学考研必背300题.docx","发展心理学")
    dp.doc_parser(r".\题库\普通心理学名词解释及论述.docx")

    dp.save_database_to_excel(r'.\example.xlsx',"发心")

    print("######################################################")
    print(dp.database)
    print("######################################################")

